#!/usr/bin/python
# -*- coding: UTF-8 -*-

import os
import sys
import pdb
import logging
from docx import Document

match_check = 0
match_total = 0
replace_total = 0

logger = logging.getLogger("mainModule")

def log_to_ui(to_ui, level, msg):
    msg_level = ""
    if "debug" == level:
        msg_level = "[DEBUG]"
        logger.debug(msg_level + msg)
    elif "error" == level:
        msg_level = "[ERROR]"
        logger.error(msg_level + msg)
    elif "info" == level:
        msg_level = "[INFO]"
        logger.info(msg_level + msg)
    else:
        logger.debug("[DEBUG]" + msg)
    if True == to_ui:
        from package_office.ui import ui_log_show
        ui_log_show(msg_level + msg)

def check_parag_match_count(para, replace_dict):
    count = 0
    for key, value in replace_dict.items():
        #strip space and \r \n
        para_new = para.text.replace(' ', '').replace('\r', '').replace('\n', '')
        if key in para_new:
            count += 1
    return count

def check_all_match_count(fdocx, replace_dict):
    count = 0
    for para in fdocx.paragraphs:
        count += check_parag_match_count(para, replace_dict)
    return count

def do_tables_replace(fdocx, replace_dict):
    if 0 == len(fdocx.tables):
        logger.info('no tables found')
        return
    cell_set = set()
    for i in range(len(fdocx.tables)):
        logger.info(f'First table:{len(fdocx.tables[i].rows)} row X {len(fdocx.tables[i].columns)} columns')
        cell_set.clear()
        for j, row in enumerate(fdocx.tables[i].rows):
            logger.info(f'row {j+1} has {len(fdocx.tables[i].columns)} columns')
            for cell in row.cells:
                if cell not in cell_set:
                    cell_set.add(cell)
                    logger.info('cell.text' + cell.text)
                    logger.info('cell.text.front' + str(cell.paragraphs))
                    for key, value in replace_dict.items():
                        if key in cell.text:
                            cell.text = cell.text.replace(key, value)

def do_re_replace(para, replace_dict):
    replace_count = 0
    for key, value in replace_dict.items():

#TODO find next
        has_find = False
        need_replace_count = 0
        total_len = 0
        pos = para.text.find(key)
        if -1 == pos:
            continue
        else:
            logger.debug ("has found,pos:%d" %pos + " arag:" + para.text)
            need_replace_count = len(key)
            for i in range(len(para.runs)):
                logger.debug ("need_replace_count %d" %need_replace_count)
                ori_len = len(para.runs[i].text)
                if (pos >= total_len) and (pos < total_len + ori_len):
                    logger.debug ("has found,ori:" + para.runs[i].text)
                    if 0 == pos and len(para.runs[i].text) <= need_replace_count:
                        para.runs[i].text = para.runs[i].text.replace(para.runs[i].text, value)
                    else:
                        para.runs[i].text = para.runs[i].text.replace(para.runs[i].text, para.runs[i].text[:pos - total_len] + value)
                    logger.debug ("has found,after:" + para.runs[i].text)
                    has_find = True
                    need_replace_count -= total_len + ori_len  - pos
                    total_len += len(para.runs[i].text)
                else:
                    if True == has_find:
                        logger.debug ("need_replace_count %d" % need_replace_count)
                        logger.debug ("len(para.runs[i].text):%d" % ori_len)
                        if need_replace_count >= ori_len:
#                            logger.debug("1replace before:" + para.runs[i].text)
                            para.runs[i].text = para.runs[i].text.replace(para.runs[i].text, "")
#                            logger.debug("2replace done:" + para.runs[i].text)
                            need_replace_count -= ori_len
                        else:
#                            logger.debug("replace before:" + para.runs[i].text)
                            para.runs[i].text = para.runs[i].text.replace(para.runs[i].text, para.runs[i].text[need_replace_count:])
                            has_find = False
                            need_replace_count = 0
                            total_len = 0
                            replace_count += 1
#                            logger.debug("replace done:" + para.runs[i].text)
                            break
                    else:
                        total_len += ori_len
                        continue
    return replace_count

def do_replace(fdocx, replace_dict):
    has_replace_count = 0
    try:
        for para in fdocx.paragraphs:
            para_match_count = check_parag_match_count(para, replace_dict)
            parag_replace_count = 0
            for i in range(len(para.runs)):
#                logger.debug (para.runs[i].text)
                for key, value in replace_dict.items():
                    if key in para.runs[i].text:
                        logger.debug ("------------------------->get " + para.runs[i].text)
                        para.runs[i].text = para.runs[i].text.replace(key, value)
                        logger.debug ("------------------------->after:" + para.runs[i].text)
                        parag_replace_count += 1
            if para_match_count != parag_replace_count:
                parag_replace_count += do_re_replace(para, replace_dict)
            has_replace_count += parag_replace_count
        do_tables_replace(fdocx, replace_dict)
    except Exception as e:
        logger.error("line:%d" % sys._getframe().f_lineno + " repr(e):", repr(e))
        return
    return has_replace_count


def do_docx(old_path, new_path, replace_dict):
    count_logical = 0
    count_real = 0
    fdocx = Document(old_path)
    count_logical = check_all_match_count(fdocx, replace_dict)
    if 0 < count_logical:
        count_real = do_replace(fdocx, replace_dict)
    if count_logical == count_real:
        if 0 < count_logical:
            log_to_ui(True, "info", "--->文件中检查到匹配项%d个，全部替换！"%count_logical)
        else:
            log_to_ui(True, "info", "--->文件中检查到匹配项%d个，跳过！"%count_logical)
    else:
        log_to_ui(True, "error", "--->文件中检查到匹配项%d"%count_logical + "个，但实际替换%d个"%count_real)
        logger.info("--->Warnning! " + old_path + " has checked same words:%d" % count_logical + " but replaced count:%d" % count_real)

    fdocx.save(new_path)
    logger.debug("--->replace done, new file " + new_path)


